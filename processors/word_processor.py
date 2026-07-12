"""
Word 文档 (.docx) 处理器

支持功能:
- 段落文本提取
- 表格内容提取
- 页眉/页脚提取
- 双栏/多栏布局检测与处理
- 文本框内容提取

双栏处理原理:
Word 的 XML 中，文本按阅读顺序连续存储（左栏从上到下 → 右栏从上到下）。
python-docx 默认按此顺序提取文本，因此对于标准的流式多栏文档，
文本顺序是正确的。但对于复杂排版（如报纸式非连续栏目），
本处理器会检测栏目结构并在元数据中标记，同时提取栏目分隔信息。
"""

import re
from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.table import Table

from .base_processor import BaseProcessor


class WordProcessor(BaseProcessor):
    processor_name = "word"
    SUPPORTED_EXTENSIONS = {".docx"}

    # Word 文档的命名空间 URI
    NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def can_handle(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext == ".docx"

    def extract_text(self, file_path: str) -> str:
        doc = Document(file_path)
        text_parts: List[str] = []

        # 1. 提取页眉/页脚
        for section in doc.sections:
            header_text = self._extract_header_footer(section.header)
            footer_text = self._extract_header_footer(section.footer)
            if header_text:
                text_parts.append(header_text)
            if footer_text:
                text_parts.append(footer_text)

        # 2. 提取正文段落
        for paragraph in doc.paragraphs:
            para_text = self._clean_text(paragraph.text)
            if para_text:
                # 检测段落样式以保留标题信息
                style_name = paragraph.style.name if paragraph.style else ""
                if self._is_heading(style_name):
                    text_parts.append(f"\n## {para_text}")
                else:
                    text_parts.append(para_text)

        # 3. 提取表格内容
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                text_parts.append(table_text)

        # 4. 提取文本框（在 XML body 中查找）
        body = doc.element.body
        textboxes = body.findall(f'.//{{{self.NS_W}}}txbxContent')
        for tb in textboxes:
            tb_text = self._extract_xml_text(tb)
            if tb_text.strip():
                text_parts.append(tb_text.strip())

        return "\n\n".join(text_parts)

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        metadata = super().extract_metadata(file_path)
        try:
            doc = Document(file_path)
            # 检测栏目布局
            column_info = self._detect_columns(doc)
            metadata.update({
                "page_count": len(doc.sections),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "column_layout": column_info,
                "has_columns": column_info["column_count"] > 1,
            })

            # 提取文档属性
            if doc.core_properties:
                cp = doc.core_properties
                metadata["title"] = cp.title or ""
                metadata["author"] = cp.author or ""
                metadata["created"] = str(cp.created) if cp.created else ""
        except Exception:
            pass
        return metadata

    # --- 内部方法 ---

    def _detect_columns(self, doc: Document) -> Dict[str, Any]:
        """
        检测文档的栏目布局。
        通过解析 sectPr → cols 元素获取栏目数和间距。
        """
        column_info = {
            "column_count": 1,
            "columns_per_section": [],
            "has_equal_width": True,
            "space_between": 0,
        }
        try:
            for i, section in enumerate(doc.sections):
                sect_pr = getattr(section, '_sectPr', None)
                if sect_pr is None:
                    continue

                cols = sect_pr.find(qn('w:cols'))
                if cols is not None:
                    num = int(cols.get(qn('w:num'), 1))
                    space = cols.get(qn('w:space'))
                    equal = cols.get(qn('w:equalWidth'), '1') == '1'

                    column_info["columns_per_section"].append(num)
                    column_info["column_count"] = max(column_info["column_count"], num)
                    column_info["has_equal_width"] = column_info["has_equal_width"] and equal
                    if space:
                        # space 单位是 twips (1/20 磅)
                        column_info["space_between"] = int(space) / 20
                else:
                    column_info["columns_per_section"].append(1)
        except Exception:
            pass
        return column_info

    def _extract_header_footer(self, hf) -> str:
        """提取页眉或页脚中的文本"""
        if hf is None:
            return ""
        parts = []
        for para in hf.paragraphs:
            text = self._clean_text(para.text)
            if text:
                parts.append(text)
        for table in hf.tables:
            table_text = self._extract_table(table)
            if table_text:
                parts.append(table_text)
        return "\n".join(parts)

    def _extract_table(self, table: Table) -> str:
        """将表格转换为可读文本"""
        rows = []
        for row in table.rows:
            cells = [self._clean_text(cell.text) for cell in row.cells]
            # 跳过全空行
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            return "[表格]\n" + "\n".join(rows) + "\n[/表格]"
        return ""

    def _extract_xml_text(self, element) -> str:
        """递归提取 XML 元素中的文本"""
        texts = []
        if element.text:
            texts.append(element.text)
        for child in element:
            texts.append(self._extract_xml_text(child))
            if child.tail:
                texts.append(child.tail)
        return " ".join(texts)

    def _is_heading(self, style_name: str) -> bool:
        """判断段落样式是否为标题"""
        return bool(re.match(r'^(Heading|heading|标题)', style_name))

    @staticmethod
    def _clean_text(text: str) -> str:
        """清理文本：去除多余空白，保留有意义的换行"""
        if not text:
            return ""
        # 合并多余空格
        text = re.sub(r' {2,}', ' ', text)
        # 去除首尾空白
        return text.strip()
